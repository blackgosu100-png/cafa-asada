import json
import re
import sys
from pathlib import Path

from docx import Document
from docx.oxml.ns import qn
from docx.shared import Pt


def set_run_font(run, name="맑은 고딕", size=10.5):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)
    run.font.size = Pt(size)


def set_paragraph_spacing(paragraph, before=0, after=4, line=1.15):
    paragraph.paragraph_format.space_before = Pt(before)
    paragraph.paragraph_format.space_after = Pt(after)
    paragraph.paragraph_format.line_spacing = line


def add_body_text(doc, text):
    for block in (text or "").split("\n"):
        paragraph = doc.add_paragraph()
        set_paragraph_spacing(paragraph)
        run = paragraph.add_run(block)
        set_run_font(run, size=10.5)


def strip_staff_checks(text):
    marker = "직원이 확인할 점"
    value = text or ""
    index = value.find(marker)
    if index == -1:
        return value.strip()
    return value[:index].strip()


def format_body_for_review(text):
    value = strip_staff_checks(text)
    value = value.replace("\r\n", "\n").replace("\r", "\n")
    value = re.sub(r"(?<!\d)([.,])(?!\d)\s*", r"\1\n\n", value)
    value = re.sub(r"\n{3,}", "\n\n", value)
    return value.strip()


def build_docx(data, output_path):
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Pt(72)
    section.bottom_margin = Pt(72)
    section.left_margin = Pt(72)
    section.right_margin = Pt(72)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "맑은 고딕"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "맑은 고딕")
    normal.font.size = Pt(10.5)

    title = data.get("editedTitle") or data.get("aiBestTitle") or "카페 검수 원고"
    paragraph = doc.add_paragraph()
    set_paragraph_spacing(paragraph, after=12)
    run = paragraph.add_run(title)
    set_run_font(run, size=10.5)

    add_body_text(doc, format_body_for_review(data.get("editedDraft") or data.get("aiDraft") or ""))

    doc.save(output_path)


if __name__ == "__main__":
    if len(sys.argv) != 3:
        raise SystemExit("usage: build_review_docx.py input.json output.docx")
    input_path = Path(sys.argv[1])
    output_path = Path(sys.argv[2])
    build_docx(json.loads(input_path.read_text(encoding="utf-8")), output_path)
